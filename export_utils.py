from docx import Document
from fpdf import FPDF
import json
from datetime import datetime
import os
import tempfile
import re
import logging

# Setup logging
logger = logging.getLogger(__name__)

class DocumentExporter:
    @staticmethod
    def sanitize_text(text):
        """
        Sanitize text to be compatible with FPDF by replacing Unicode characters 
        with their closest ASCII equivalents
        """
        # Define replacements for common Unicode characters
        replacements = {
            '\u2014': '--',  # em dash
            '\u2013': '-',   # en dash
            '\u2018': "'",   # left single quotation mark
            '\u2019': "'",   # right single quotation mark
            '\u201c': '"',   # left double quotation mark
            '\u201d': '"',   # right double quotation mark
            '\u2022': '*',   # bullet
            '\u2026': '...', # ellipsis
            '\u00a9': '(c)', # copyright symbol
            '\u00ae': '(R)', # registered trademark symbol
            '\u2122': '(TM)',# trademark symbol
            # Add more replacements as needed
        }
        
        # Replace unsupported characters
        for unicode_char, replacement in replacements.items():
            text = text.replace(unicode_char, replacement)
        
        # Remove remaining non-Latin-1 characters
        text = re.sub(r'[^\x00-\xff]', ' ', text)
        
        return text
    
    @staticmethod
    def export_pdf(documentation, filename=None, output_dir=None):
        """
        Export documentation to PDF format
        
        Args:
            documentation (str): Documentation text to export
            filename (str, optional): Output filename. Defaults to None.
            output_dir (str, optional): Directory to save the file. Defaults to None.
            
        Returns:
            str: Path to the generated PDF file
        """
        try:
            # Sanitize the documentation text
            sanitized_documentation = DocumentExporter.sanitize_text(documentation)
            
            # If filename is not provided, generate one using current date and time
            if filename is None:
                filename = f"documentation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            
            # If output_dir is not provided, use 'exports' or temp directory
            if output_dir is None:
                output_dir = os.path.join(tempfile.gettempdir(), 'exports')
            
            # Ensure the output directory exists
            os.makedirs(output_dir, exist_ok=True)
            
            # Set output file path
            output_path = os.path.join(output_dir, filename)
            
            # Create a PDF instance
            pdf = FPDF()
            pdf.add_page()
            
            # Set font for the PDF
            pdf.set_font("Arial", size=12)
            
            # Split the documentation into lines and add to PDF
            for line in sanitized_documentation.split('\n'):
                # Limit line length to prevent overflows
                chunks = [line[i:i+80] for i in range(0, len(line), 80)]
                for chunk in chunks:
                    pdf.cell(0, 10, txt=chunk, ln=True)
            
            # Output the PDF to the specified path
            pdf.output(output_path)
            
            logger.info(f"PDF generated successfully at: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to generate PDF: {str(e)}", exc_info=True)
            raise
    
    @staticmethod
    def export_docx(documentation, filename=None, output_dir=None):
        """
        Export documentation to DOCX format
        
        Args:
            documentation (str): Documentation text to export
            filename (str, optional): Output filename. Defaults to None.
            output_dir (str, optional): Directory to save the file. Defaults to None.
            
        Returns:
            str: Path to the generated DOCX file
        """
        try:
            # If filename is not provided, generate one using current date and time
            if filename is None:
                filename = f"documentation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            
            # If output_dir is not provided, use 'exports' or temp directory
            if output_dir is None:
                output_dir = os.path.join(tempfile.gettempdir(), 'exports')
            
            # Ensure the output directory exists
            os.makedirs(output_dir, exist_ok=True)
            
            # Set output file path
            output_path = os.path.join(output_dir, filename)
            
            # Create a Document instance
            doc = Document()
            doc.add_heading('Code Documentation', 0)
            
            # Add content to document
            for paragraph in documentation.split('\n\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            
            # Save the document
            doc.save(output_path)
            
            logger.info(f"DOCX generated successfully at: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to generate DOCX: {str(e)}", exc_info=True)
            raise

# Example usage
if __name__ == "__main__":
    doc = "This is an example documentation.\nIt contains multiple lines.\nEach line will be written to the PDF."
    output_file = DocumentExporter.export_pdf(doc)
    print(f"PDF generated: {output_file}")