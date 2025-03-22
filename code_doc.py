import streamlit as st
import ollama
import ast
import json
from datetime import datetime
import docx
from fpdf import FPDF
import git
import os
from pathlib import Path
import hashlib
from typing import Optional, Dict, List

# Authentication and User Management
class UserManager:
    def __init__(self, data_path: str = "users.json"):
        self.data_path = data_path
        self._load_users()

    def _load_users(self):
        if os.path.exists(self.data_path):
            with open(self.data_path, 'r') as f:
                self.users = json.load(f)
        else:
            self.users = {}
            self._save_users()

    def _save_users(self):
        with open(self.data_path, 'w') as f:
            json.dump(self.users, f)

    def register_user(self, username: str, password: str) -> bool:
        if username in self.users:
            return False
        self.users[username] = {
            'password_hash': hashlib.sha256(password.encode()).hexdigest(),
            'created_at': datetime.now().isoformat()
        }
        self._save_users()
        return True

    def verify_user(self, username: str, password: str) -> bool:
        if username not in self.users:
            return False
        return self.users[username]['password_hash'] == hashlib.sha256(password.encode()).hexdigest()

# Documentation History Manager
class DocumentationHistory:
    def __init__(self, history_path: str = "doc_history"):
        self.history_path = Path(history_path)
        self.history_path.mkdir(exist_ok=True)

    def save_documentation(self, username: str, code: str, documentation: str) -> str:
        timestamp = datetime.now().isoformat()
        doc_id = hashlib.md5(f"{username}{timestamp}".encode()).hexdigest()
        
        user_path = self.history_path / username
        user_path.mkdir(exist_ok=True)
        
        with open(user_path / f"{doc_id}.json", 'w') as f:
            json.dump({
                'timestamp': timestamp,
                'code': code,
                'documentation': documentation
            }, f)
        
        return doc_id

    def get_user_history(self, username: str) -> List[Dict]:
        user_path = self.history_path / username
        if not user_path.exists():
            return []
        
        history = []
        for file in user_path.glob("*.json"):
            with open(file, 'r') as f:
                data = json.load(f)
                data['id'] = file.stem
                history.append(data)
        
        return sorted(history, key=lambda x: x['timestamp'], reverse=True)

# File Export Functions
def export_to_pdf(documentation: str, filename: str):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    # Split text into lines and add to PDF
    for line in documentation.split('\n'):
        pdf.multi_cell(0, 10, txt=line)
    
    pdf.output(filename)

def export_to_docx(documentation: str, filename: str):
    doc = docx.Document()
    doc.add_heading('Code Documentation', 0)
    
    for paragraph in documentation.split('\n\n'):
        doc.add_paragraph(paragraph)
    
    doc.save(filename)

# Version Control Integration
class GitManager:
    def __init__(self, repo_path: str = "code_repository"):
        self.repo_path = Path(repo_path)
        self._initialize_repo()

    def _initialize_repo(self):
        if not (self.repo_path / '.git').exists():
            self.repo_path.mkdir(exist_ok=True)
            git.Repo.init(self.repo_path)
        self.repo = git.Repo(self.repo_path)

    def save_code_version(self, username: str, code: str, message: str) -> str:
        file_path = self.repo_path / f"{username}_latest.py"
        with open(file_path, 'w') as f:
            f.write(code)
        
        self.repo.index.add([str(file_path)])
        commit = self.repo.index.commit(f"{username}: {message}")
        return commit.hexsha

# Enhanced Code Analysis
def analyze_code_structure(code: str) -> tuple:
    try:
        tree = ast.parse(code)
        
        functions = []
        classes = []
        imports = []
        
        for node in ast.walk(tree):
            if isinstance(node, ast.FunctionDef):
                functions.append({
                    'name': node.name,
                    'args': [arg.arg for arg in node.args.args],
                    'docstring': ast.get_docstring(node)
                })
            elif isinstance(node, ast.ClassDef):
                classes.append({
                    'name': node.name,
                    'bases': [base.id for base in node.bases if isinstance(base, ast.Name)],
                    'docstring': ast.get_docstring(node)
                })
            elif isinstance(node, ast.Import):
                imports.extend(alias.name for alias in node.names)
            elif isinstance(node, ast.ImportFrom):
                imports.append(f"{node.module}.{node.names[0].name}")
        
        return functions, classes, imports
    except Exception as e:
        return [], [], []

# Error Handler
class ErrorHandler:
    @staticmethod
    def handle_error(func):
        def wrapper(*args, **kwargs):
            try:
                return func(*args, **kwargs)
            except Exception as e:
                st.error(f"Error in {func.__name__}: {str(e)}")
                return None
        return wrapper

# Initialize managers
user_manager = UserManager()
history_manager = DocumentationHistory()
git_manager = GitManager()

# Streamlit UI Components
def render_login_ui():
    st.sidebar.header("Authentication")
    
    auth_option = st.sidebar.radio("Choose Option", ["Login", "Register"])
    
    username = st.sidebar.text_input("Username")
    password = st.sidebar.text_input("Password", type="password")
    
    if auth_option == "Login":
        if st.sidebar.button("Login"):
            if user_manager.verify_user(username, password):
                st.session_state['user'] = username
                st.success("Login successful!")
            else:
                st.error("Invalid credentials!")
    else:
        if st.sidebar.button("Register"):
            if user_manager.register_user(username, password):
                st.success("Registration successful! Please login.")
            else:
                st.error("Username already exists!")

def generate_documentation(code_input: str) -> str:
    """Generate documentation using Ollama"""
    prompt = f"""
    Analyze this Python code and generate comprehensive documentation:
    {code_input}
    
    Include:
    1. Overview of the code's purpose
    2. Functions and classes documentation
    3. Dependencies and requirements
    4. Usage examples
    
    Format the response in markdown.
    """
    
    response = ollama.chat(
        model="deepseek-r1:1.5b",
        messages=[{"role": "user", "content": prompt}]
    )
    
    return response['message']['content']

def render_main_ui():
    st.title("Enhanced GPT-Based Documentation Generator")
    
    # Code input
    code_input = st.text_area("Paste your Python code here:", height=300)
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Generate Documentation"):
            if code_input.strip():
                # Analyze code
                functions, classes, imports = analyze_code_structure(code_input)
                
                # Generate documentation
                doc_output = generate_documentation(code_input)
                
                # Save to history
                doc_id = history_manager.save_documentation(
                    st.session_state['user'],
                    code_input,
                    doc_output
                )
                
                # Save to version control
                commit_hash = git_manager.save_code_version(
                    st.session_state['user'],
                    code_input,
                    "Updated documentation"
                )
                
                # Display results
                st.write("### Code Analysis:")
                st.json({
                    'functions': functions,
                    'classes': classes,
                    'imports': imports
                })
                
                st.write("### Generated Documentation:")
                st.markdown(doc_output)
                
                # Export options
                st.write("### Export Documentation")
                export_format = st.selectbox("Choose format:", ["PDF", "DOCX"])
                if st.button("Download"):
                    if export_format == "PDF":
                        export_to_pdf(doc_output, "documentation.pdf")
                        st.download_button(
                            "Download PDF",
                            open("documentation.pdf", "rb"),
                            "documentation.pdf",
                            "application/pdf"
                        )
                    else:
                        export_to_docx(doc_output, "documentation.docx")
                        st.download_button(
                            "Download DOCX",
                            open("documentation.docx", "rb"),
                            "documentation.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

def render_history_ui():
    st.sidebar.header("Documentation History")
    history = history_manager.get_user_history(st.session_state['user'])
    
    for entry in history:
        with st.sidebar.expander(f"Doc {entry['timestamp'][:10]}"):
            st.code(entry['code'], language="python")
            st.markdown(entry['documentation'])

# Main app
def main():
    # Initialize session state
    if 'user' not in st.session_state:
        st.session_state['user'] = None
    
    # Custom CSS for better UI
    st.markdown("""
        <style>
        .stButton>button {
            width: 100%;
        }
        .stTextArea>div>div>textarea {
            background-color: #f0f2f6;
        }
        </style>
    """, unsafe_allow_html=True)
    
    if st.session_state['user'] is None:
        render_login_ui()
    else:
        render_main_ui()
        render_history_ui()

if __name__ == "__main__":
    main()